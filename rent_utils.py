import re
import pdfplumber
from docx import Document
from datetime import datetime
import os
import threading
from groq import Groq

GROQ_API_KEY = os.getenv("GROQ_API_KEY")
if not GROQ_API_KEY:
    raise ValueError("No GROQ_API_KEY environment variable has been set.")


client = Groq(
    api_key=GROQ_API_KEY
)

def translate_text(text):
    if not text:  # Skip translation if text is empty
        return ""
    
    prompt = f"""Translate the following Swedish text to English, do not add any commentary to the translation:
    {text}
    Translation:"""
    
    chat_completion = client.chat.completions.create(
        messages=[
            {
                "role": "user",
                "content": prompt,
            }
        ],
        model="llama3-8b-8192",
        temperature=0.1,
        max_tokens=200,
        top_p=1,
    )
    
    response = chat_completion.choices[0].message.content.strip()
    
    # Extract only the translated text
    translation_start = response.find('"') + 1
    translation_end = response.rfind('"')
    if translation_start > 0 and translation_end > translation_start:
        return response[translation_start:translation_end]
    else:
        return response  # Return full response if we can't extract the translation

def schedule_file_deletion(file_path, delay_seconds):
    def delete_file():
        print(f"Attempting to delete file: {file_path}")
        if os.path.exists(file_path):
            os.remove(file_path)
            print(f"Deleted file: {file_path}")
        else:
            print(f"File not found for deletion: {file_path}")

    print(f"Scheduling deletion of {file_path} in {delay_seconds} seconds")
    threading.Timer(delay_seconds, delete_file).start()

def extract_info_from_pdf(pdf_path):
    print(f"Extracting info from PDF: {pdf_path}")
    with pdfplumber.open(pdf_path) as pdf:
        if not pdf.pages:
            raise ValueError("PDF has no pages.")
        
        first_page_text = pdf.pages[0].extract_text()
        if not first_page_text:
            first_page_text = "" # Ensure it's a string if no text extracted
        
        print(f"Extracted text from first page: {first_page_text}")

        # Regex to find signees like (1) Name, (2) Another Name
        # It captures the name part after the (number) prefix.
        signee_matches = re.finditer(r'\(\d+\)\s*([^,\(]+)', first_page_text)
        signees = [match.group(1).strip() for match in signee_matches]
        
        if not signees:
            print("Warning: No signees found with the pattern '(number) Name' on the first page. The generated document might be missing signee information.")

        address_match = re.search(r'adress\s+(.*?),', first_page_text, re.IGNORECASE)
        address = address_match.group(1).strip() if address_match else "N/A"
        
        transaction_search = re.search(r'Transaktion\s+(\S+)', first_page_text)
        transaction_id = transaction_search.group(1).strip() if transaction_search else "N/A"

        current_rent = None
        # Search for current_rent across all pages
        for i, page in enumerate(pdf.pages):
            page_text = page.extract_text()
            if page_text: # Ensure text was extracted
                # Try primary pattern
                match = re.search(r'Hyran är\s+([\d\s]+)', page_text)
                if match:
                    current_rent = match.group(1).strip().replace(" ", "")
                    break
                # Try alternative pattern if primary not found on this page
                match_alt = re.search(r'nuvarande hyra[:\s]+([\d\s]+)kr', page_text, re.IGNORECASE)
                if match_alt:
                    current_rent = match_alt.group(1).strip().replace(" ", "")
                    break
        
        if current_rent is None:
            print("Warning: Current rent not found in the PDF using known patterns. Defaulting to 'N/A' or consider raising an error.")
            current_rent = "N/A" # Or raise ValueError("Current rent not found in the PDF.")

        print(f"Extracted info: Signees: {signees}, Address: {address}, Transaction ID: {transaction_id}, Current Rent: {current_rent}")
        return signees, address, transaction_id, current_rent

def replace_placeholders(doc, placeholders):
    def process_paragraph(paragraph):
        text = paragraph.text
        
        # First, replace all known placeholders with their values
        for key, value in placeholders.items():
            text = text.replace(key, str(value))
        
        # Then, remove any [SIGNEE_X] placeholders that were not filled
        # because there weren't that many signees.
        # Iterate up to a reasonable max, e.g., 20 potential signees.
        for i in range(1, 21): # Check for [SIGNEE_1] up to [SIGNEE_20]
            leftover_placeholder = f'[SIGNEE_{i}]'
            # If the placeholder is still literally in the text after the first loop, 
            # it means it was NOT a key in 'placeholders' that got replaced with actual content.
            if leftover_placeholder in text:
                 text = text.replace(leftover_placeholder, '')

        paragraph.clear()
        
        # Determine if the line should be bold
        # A line is considered for bolding if it starts with (number) or is a signature line.
        is_bold_line = bool(re.match(r'^\s*\(\d+\)', text)) or text.strip().startswith("_____")
        
        # Only add run if there's non-whitespace text to avoid empty runs
        if text.strip(): 
            run = paragraph.add_run(text)
            run.bold = is_bold_line
        # If text becomes empty or only whitespace after replacements, the paragraph will effectively be cleared.

    for paragraph in doc.paragraphs:
        process_paragraph(paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_paragraph(paragraph)

def set_font_to_times_new_roman(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'

def create_new_rent_increase_docx(template_path, signees, application_date, current_rent, new_rent, service_fee, address, transaction_id, free_text, end_date, output_path):
    print(f"Creating new rent increase DOCX: {output_path} for signees: {signees}")
    doc = Document(template_path)

    rounded_service_fee = round(service_fee)

    if end_date:
        when_se = f"till och med {end_date.strftime('%Y-%m-%d')}. Därefter återgår hyran till föregående belopp"
        when_en = f"{end_date.strftime('%Y-%m-%d')}. After which the rent returns to the previous amount"
    else:
        when_se = "tillsvidare"
        when_en = "further notice"

    print(f"Original free_text: {free_text}")
    free_text_en = translate_text(free_text) if free_text else ''
    print(f"Translated free_text: {free_text_en}")

    placeholders = {
        '[ADDRESS]': address,
        '[TRANSACTION_ID]': transaction_id,
        '[CURRENT_RENT]': current_rent,
        '[NEW_RENT]': f"{new_rent:.0f}",
        '[SERVICE_FEE]': str(rounded_service_fee),
        '[APPLICATION_DATE]': application_date.strftime('%Y-%m-%d'),
        '[TODAYS_DATE]': datetime.today().strftime('%Y-%m-%d'),
        '[FREE_TEXT]': free_text,
        '[FREE_TEXT_EN]': free_text_en,
        '[WHEN_SE]': when_se,
        '[WHEN_EN]': when_en,
    }

    # Add dynamic signee placeholders with full formatting
    if signees: # Check if the signees list is not empty
        # Format landlord (SIGNEE_1)
        placeholders['[SIGNEE_1]'] = f"(1) {signees[0]}, +46 xxx xx xx (Hyresvärden) och"
        
        # Format tenants (SIGNEE_2, SIGNEE_3, ...)
        for i in range(1, len(signees)):
            tenant_name = signees[i]
            placeholder_key = f'[SIGNEE_{i+1}]' # e.g., [SIGNEE_2], [SIGNEE_3]
            placeholder_value = f"({i+1}) {tenant_name}, +46 xxx xx xx (Hyresgästen)"
            placeholders[placeholder_key] = placeholder_value
    else:
        print("Warning: No signees provided to create_new_rent_increase_docx. Signee placeholders in the document might not be filled.")

    print("Placeholders to be replaced:", placeholders)

    replace_placeholders(doc, placeholders)
    set_font_to_times_new_roman(doc)

    doc.save(output_path)
    print(f"DOCX saved: {output_path}")

    return output_path